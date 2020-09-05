import docx
from datetime import date

doc = docx.Document()


def takeInputForTable():
    x = True
    count = 1
    records = []
    while x:
        particular = input("Enter The Particular :")
        quantity = float(input("Enter The Quantity :"))
        hsn = input("Enter The HSN Code :")
        rate = float(input("Enter The Rate :"))
        total = rate * quantity
        t = Table(particular,quantity,hsn,total)
        print(total)
        item = {
            "sno": count,
            "particular" : particular,
            "quantity" : quantity,
            "rate" : rate,
             "hsn" : hsn,
             "total" : t.Total
        }
        records.append(item)
        print("Item Added !!!")
        print("---------")
        print("---------")
        check = input("Enter More records : Y/N:")
        if check == "N" or check == "n":
            x = False
            return records
        count+=1

class Table:
    def __init__(self,particular,qty, Rate, Total):
        self.particular = particular
        self.qty = qty
        self.HSNCode = 995469
        self.Rate = Rate
        self.Total = Total

    
            

class PapaQuotation(Table):
    def __init__(self,date,recipient, address1, address2,subject,note):
        self.date = date
        self.recipient = recipient
        self.address1 = address1
        self.address2 = address2
        self.subject = subject
        self.note = note


    def systemOut(self):
        try:
            count = 1
            doc.add_paragraph("To								 Date-" + self.date)
            doc.add_paragraph("The Manager")
            doc.add_paragraph(self.recipient)
            doc.add_paragraph(self.address1)
            doc.add_paragraph(self.address2)
            doc.add_paragraph("Subject : " +self.subject)
            doc.add_paragraph("Dear Sir,")
            records = takeInputForTable()
            table = doc.add_table(rows = 1, cols = 6)
            hdr_cells = table.rows[0].cells
            
            
            hdr_cells[0].text = "Sno"
            hdr_cells[1].text = "Particular"
            hdr_cells[2].text = "QTY"
            hdr_cells[3].text = "HSN Code"
            hdr_cells[4].text = "Rate"
            hdr_cells[5].text = "Total"
            
            for record in records:
                row_cells = table.add_row().cells
                row_cells[0].text =str(record.get("sno"))
                row_cells[1].text =str(record.get("particular"))
                row_cells[2].text =str(record.get("quantity"))
                row_cells[3].text =str(record.get("hsn"))
                row_cells[4].text = str(record.get("rate"))
                row_cells[5].text =str(record.get("total"))

            for item in self.note:
                doc.add_paragraph(count + "."+ item)
                count+=1
        except :
            pass
        finally:
            name = "quotation/{}-quotation.docx".format(str(date.today()))
            
            doc.save(name)
        
def takeInput():
    date = input("Enter The Date :")
    recipient = input("Enter The Recipient :")
    address1 = input("Enter The Address 1:")
    address2 = input("Enter The Address 2:")
    subject = input("Enter The Subject:")
    note = input("Enter The Note:").split("#")
    papa= PapaQuotation(date,recipient,address1,address2,subject,note)
    papa.systemOut()

takeInput()



    